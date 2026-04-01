import streamlit as st
import io
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("Kripya 'python-docx' library install karein: pip install python-docx")

st.title("📝 GST NOC / Consent Letter Generator")
st.write("Property Owner aur Applicant ki details fill karein aur 1-click mein GST portal ke liye ready-to-print NOC Word Document (.docx) download karein.")

# ==========================================
# 🚀 NOC FORM UI
# ==========================================
with st.form("noc_form"):
    st.subheader("👤 Owner & Applicant Details")
    
    col1, col2 = st.columns(2)
    with col1:
        owner_name = st.text_input("Property Owner Name (e.g., Ramesh Kumar)")
        applicant_name = st.text_input("Applicant / Business Name (e.g., Suresh Kumar / Freyora Textiles)")
    with col2:
        relation = st.selectbox("Owner's Relation with Applicant", ["Father", "Mother", "Spouse", "Son", "Daughter", "Brother", "Landlord", "Other"])
        if relation == "Other":
            relation = st.text_input("Please specify relation")
            
    st.subheader("📍 Property Details")
    property_address = st.text_area("Full Address of the Premises (As per Electricity Bill)")
    
    col3, col4 = st.columns(2)
    with col3:
        place = st.text_input("Place (e.g., Surat)")
    with col4:
        date_of_noc = st.date_input("Date of Issue", value=datetime.today())

    submit_btn = st.form_submit_button("📄 Generate NOC Document", type="primary")

# ==========================================
# 💻 WORD DOCX GENERATION LOGIC
# ==========================================
if submit_btn:
    if not owner_name or not applicant_name or not property_address:
        st.warning("⚠️ Kripya Owner Name, Applicant Name aur Address zaroor fill karein.")
    else:
        with st.spinner("Word Document ban raha hai..."):
            # Naya Word Document banayein
            doc = Document()
            
            # Heading
            heading = doc.add_heading('CONSENT LETTER / NO OBJECTION CERTIFICATE', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph() # Khali line
            
            # "To Whomsoever" Title (Center aligned)
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.add_run("TO WHOMSOEVER IT MAY CONCERN").bold = True
            
            doc.add_paragraph() # Khali line
            
            # Paragraph 1
            p1 = doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p1.add_run("This is to certify that I, ")
            p1.add_run(f"{owner_name.upper()}").bold = True
            p1.add_run(", am the legal owner of the property situated at ")
            p1.add_run(f"{property_address}").bold = True
            p1.add_run(".")
            
            doc.add_paragraph() # Khali line
            
            # Paragraph 2
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p2.add_run(f"I do hereby solemnly affirm and declare that I have no objection if my {relation.lower()}, ")
            p2.add_run(f"{applicant_name.upper()}").bold = True
            p2.add_run(", uses the above-mentioned premises as the Principal Place of Business / Additional Place of Business for obtaining Goods and Services Tax (GST) Registration and operating the business.")
            
            doc.add_paragraph() # Khali line
            
            # Paragraph 3
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p3.add_run("I further declare that this consent is being issued out of my own free will and without any pressure or coercion.")
            
            doc.add_paragraph()
            doc.add_paragraph() # Do khali lines signature se pehle
            
            # Signatures and Date
            sign_para = doc.add_paragraph()
            sign_para.add_run(f"Date: {date_of_noc.strftime('%d-%m-%Y')}\n")
            sign_para.add_run(f"Place: {place}\n\n\n\n")
            
            sign_para.add_run("_________________________\n")
            sign_para.add_run("(Signature of the Owner)\n")
            sign_para.add_run(f"Name: {owner_name.upper()}")

            # Save in memory
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            st.success("✅ NOC Document successfully generated!")
            
            # Download Button
            st.download_button(
                label="📥 Download Word File (.docx)",
                data=doc_buffer,
                file_name=f"NOC_{applicant_name.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
